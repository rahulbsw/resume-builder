#!/usr/bin/env python3
"""Build polished resume DOCX files from structured resume JSON.

Usage:
  python build_resume_docx.py --input resume.json --output ats-resume.docx --profile ats-safe
  python build_resume_docx.py --input resume.json --output designed-resume.docx --profile modern-technical

If --input is omitted, the script generates a placeholder template.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "1F4D78"
ACCENT_BRIGHT = "2563EB"
INK = "172033"
MUTED = "5D667A"
LINE = "D9DEE8"
SOFT = "F4F7FB"


DEFAULT_DATA: dict[str, Any] = {
    "name": "Candidate Name",
    "target_role": "Staff Platform Engineer",
    "contact": {
        "email": "email@example.com",
        "phone": "+1 555 010 0000",
        "location": "City, ST",
        "linkedin": "linkedin.com/in/profile",
        "github": "github.com/profile",
    },
    "summary": (
        "Evidence-backed summary focused on platform engineering, distributed systems, "
        "cloud infrastructure, reliability, and cross-functional leadership."
    ),
    "skills": {
        "Languages": ["Java", "Python", "Go"],
        "Cloud and Platform": ["Kubernetes", "Terraform", "AWS", "Azure"],
        "Architecture": ["Microservices", "Distributed Systems", "API Platforms"],
        "Operations": ["Observability", "Incident Management", "CI/CD"],
    },
    "experience": [
        {
            "company": "Company Name",
            "title": "Staff Platform Engineer",
            "dates": "2021 - Present",
            "bullets": [
                "Led platform modernization for a distributed service ecosystem, improving reliability and developer productivity across multiple engineering teams.",
                "Designed cloud infrastructure and deployment automation using Kubernetes, Terraform, and CI/CD practices while reducing operational risk.",
                "Partnered with product, security, and operations stakeholders to deliver customer-facing capabilities with measurable reliability improvements.",
            ],
        },
        {
            "company": "Previous Company",
            "title": "Senior Software Engineer",
            "dates": "2018 - 2021",
            "bullets": [
                "Built backend services and APIs supporting high-availability product workflows.",
                "Improved observability and incident response through standardized metrics, logging, and service ownership practices.",
            ],
        },
    ],
    "projects": [
        {
            "name": "Platform Reliability Program",
            "bullets": [
                "Established service health standards, runbooks, and dashboards for critical systems.",
                "Reduced incident triage time by improving telemetry, ownership boundaries, and operational documentation.",
            ],
        }
    ],
    "open_source": [
        "Contributed bug fixes, reviews, or documentation to public repositories relevant to cloud-native engineering."
    ],
    "education": ["Degree, School"],
    "certifications": ["Certification Name"],
}


def hex_color(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color: str = LINE, size: str = "8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    font = run.font
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = hex_color(color)
    if bold is not None:
        font.bold = bold


def load_data(path: str | None) -> dict[str, Any]:
    if not path:
        return DEFAULT_DATA
    with open(path, "r", encoding="utf-8") as handle:
        data = json.load(handle)
    merged = DEFAULT_DATA.copy()
    merged.update(data)
    return merged


def setup_doc(profile: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    margin = 0.62 if profile != "ats-safe" else 0.72
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = hex_color(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.color.rgb = hex_color(ACCENT if profile != "ats-safe" else INK)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(13.5)
    styles["Heading 2"].font.size = Pt(11.5)
    styles["Heading 3"].font.size = Pt(10.5)
    return doc


def add_header(doc: Document, data: dict[str, Any], profile: str) -> None:
    name = data.get("name", "Candidate Name")
    target = data.get("target_role", "")
    contact = data.get("contact", {})
    contact_line = " | ".join(
        str(contact.get(key, "")).strip()
        for key in ["email", "phone", "location", "linkedin", "github"]
        if str(contact.get(key, "")).strip()
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if profile == "ats-safe" else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(name)
    set_run(r, size=25 if profile != "ats-safe" else 22, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(1)

    if target:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if profile == "ats-safe" else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(target)
        set_run(r, size=11, color=ACCENT if profile != "ats-safe" else INK, bold=True)
        p.paragraph_format.space_after = Pt(2)

    if contact_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if profile == "ats-safe" else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(contact_line)
        set_run(r, size=9.5, color=MUTED)
        p.paragraph_format.space_after = Pt(8)

    if profile != "ats-safe":
        add_bottom_border(p, color=ACCENT_BRIGHT, size="12")


def add_section(doc: Document, title: str, profile: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title.upper() if profile != "ats-safe" else title)
    set_run(r, size=10.5 if profile != "ats-safe" else 11, color=ACCENT if profile != "ats-safe" else INK, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    if profile != "ats-safe":
        add_bottom_border(p, color=LINE, size="4")


def add_summary(doc: Document, data: dict[str, Any], profile: str) -> None:
    add_section(doc, "Summary", profile)
    summary = data.get("summary", "")
    if profile == "ats-safe":
        doc.add_paragraph(summary)
        return
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(summary)
    set_run(r, size=10.5, color=INK)


def add_skills(doc: Document, data: dict[str, Any], profile: str) -> None:
    skills = data.get("skills", {})
    add_section(doc, "Technical Skills", profile)
    if profile == "ats-safe":
        for category, values in skills.items():
            p = doc.add_paragraph()
            r = p.add_run(f"{category}: ")
            set_run(r, bold=True)
            p.add_run(", ".join(values))
        return

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for category, values in skills.items():
        cells = table.add_row().cells
        cells[0].width = Inches(1.45)
        cells[1].width = Inches(5.65)
        for cell in cells:
            set_cell_margins(cell, top=80, bottom=80, start=110, end=110)
        set_cell_shading(cells[0], "E8EEF5")
        cells[0].paragraphs[0].add_run(category).bold = True
        cells[1].paragraphs[0].add_run(", ".join(values))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def add_experience(doc: Document, data: dict[str, Any], profile: str) -> None:
    add_section(doc, "Professional Experience", profile)
    for role in data.get("experience", []):
        p = doc.add_paragraph()
        title = f"{role.get('company', '')} | {role.get('title', '')}"
        r = p.add_run(title.strip(" |"))
        set_run(r, size=10.5, color=INK, bold=True)
        dates = role.get("dates")
        if dates:
            r = p.add_run(f"    {dates}")
            set_run(r, size=9.5, color=MUTED)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        for bullet in role.get("bullets", []):
            add_bullet(doc, bullet)


def add_projects(doc: Document, data: dict[str, Any], profile: str) -> None:
    projects = data.get("projects", [])
    if not projects:
        return
    add_section(doc, "Selected Projects", profile)
    for project in projects:
        p = doc.add_paragraph()
        r = p.add_run(project.get("name", "Project Name"))
        set_run(r, size=10.5, color=INK, bold=True)
        for bullet in project.get("bullets", []):
            add_bullet(doc, bullet)


def add_simple_list_section(doc: Document, title: str, values: list[str], profile: str) -> None:
    if not values:
        return
    add_section(doc, title, profile)
    for value in values:
        add_bullet(doc, value)


def add_footer_note(doc: Document, profile: str) -> None:
    if profile == "ats-safe":
        return
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Designed resume companion - use ATS-safe version for job portals")
    set_run(run, size=8, color=MUTED)


def build_docx(data: dict[str, Any], output: Path, profile: str) -> None:
    doc = setup_doc(profile)
    add_header(doc, data, profile)
    add_summary(doc, data, profile)
    add_skills(doc, data, profile)
    add_experience(doc, data, profile)
    add_projects(doc, data, profile)
    add_simple_list_section(doc, "Open Source", data.get("open_source", []), profile)
    add_simple_list_section(doc, "Education", data.get("education", []), profile)
    add_simple_list_section(doc, "Certifications", data.get("certifications", []), profile)
    add_footer_note(doc, profile)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build resume DOCX from structured JSON.")
    parser.add_argument("--input", help="Resume JSON input. Uses placeholder data when omitted.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    parser.add_argument(
        "--profile",
        default="modern-technical",
        choices=["ats-safe", "modern-technical", "executive-compact"],
        help="Template profile.",
    )
    args = parser.parse_args()

    data = load_data(args.input)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    build_docx(data, output, args.profile)


if __name__ == "__main__":
    main()
